"""build_docx_s91.py
Build supplementary section S9.1 (Kernel Rank) as a Word .docx with native Word
equations (OMML) and embedded figures.

Equations: LaTeX -> MathML (latex2mathml) -> OMML (Office MML2OMML.XSL) ->
inserted into the docx XML so Word treats them as real equation objects.
"""
import os
from lxml import etree
import latex2mathml.converter as l2m
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

HERE = os.path.dirname(os.path.abspath(__file__))
FIGDIR = os.path.join(HERE, "figures")
REPO = os.path.abspath(os.path.join(HERE, "..", "..", "..", ".."))
OUT = os.path.join(REPO, "SI_S9_1_Kernel_Rank.docx")
XSL = r"C:\Program Files\Microsoft Office\root\Office16\MML2OMML.XSL"

_transform = etree.XSLT(etree.parse(XSL))

_MNS = "{http://schemas.openxmlformats.org/officeDocument/2006/math}"

def _fix_nary(root):
    """MML2OMML.XSL leaves the n-ary (sum) operand slot <m:e> empty and emits
    the summand as a following sibling -> Word draws an empty box. Move the
    summand into <m:e> so the sum renders correctly."""
    for nary in list(root.iter(_MNS + "nary")):
        e = nary.find(_MNS + "e")
        if e is None:
            e = etree.SubElement(nary, _MNS + "e")
        if len(e) == 0:
            nxt = nary.getnext()
            if nxt is not None:
                nary.getparent().remove(nxt)
                e.append(nxt)
    return root

def omml(latex):
    """LaTeX string -> <m:oMath> lxml element (native Word equation)."""
    mathml = l2m.convert(latex)
    root = _transform(etree.fromstring(mathml.encode("utf-8"))).getroot()
    return _fix_nary(root)

# MML2OMML.XSL renders \begin{bmatrix} as a bracket-less equation array; the
# feature matrix is hand-built as a proper OMML matrix (m:m) inside a
# stretchy delimiter (m:d) so the square brackets size to the stack.
_M = "http://schemas.openxmlformats.org/officeDocument/2006/math"

def _psi_row(sub):
    return (f'<m:mr><m:e><m:sSubSup><m:e><m:r><m:t>ψ</m:t></m:r></m:e>'
            f'<m:sub><m:r><m:t>{sub}</m:t></m:r></m:sub>'
            f'<m:sup><m:r><m:t>⊤</m:t></m:r></m:sup>'
            f'</m:sSubSup></m:e></m:mr>')

EQ1_XML = (
    f'<m:oMath xmlns:m="{_M}">'
    '<m:r><m:t>Ψ = </m:t></m:r>'
    '<m:d><m:dPr><m:begChr m:val="["/><m:endChr m:val="]"/></m:dPr><m:e><m:m>'
    + _psi_row("1") + _psi_row("2")
    + '<m:mr><m:e><m:r><m:t>⋮</m:t></m:r></m:e></m:mr>'
    + _psi_row("N")
    + '</m:m></m:e></m:d>'
    '<m:r><m:t> ∈ </m:t></m:r>'
    '<m:sSup><m:e><m:r><m:t>ℝ</m:t></m:r></m:e>'
    '<m:sup><m:r><m:t>N×D</m:t></m:r></m:sup></m:sSup>'
    '</m:oMath>')

# ----------------------------------------------------------------- doc setup
doc = Document()
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)

def body(segs, justify=True):
    """Add a paragraph from segments: (text, flags) with flags in {i,b,s,p}."""
    p = doc.add_paragraph()
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for text, fl in segs:
        r = p.add_run(text)
        r.italic = "i" in fl
        r.bold = "b" in fl
        if "s" in fl:
            r.font.subscript = True
        if "p" in fl:
            r.font.superscript = True
    return p

def equation(latex):
    """Add a centred display equation as a native Word equation."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p._p.append(omml(latex))
    return p

def equation_raw(xml):
    """Add a centred display equation from a raw OMML string."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p._p.append(etree.fromstring(xml.encode("utf-8")))
    return p

def caption(segs):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(10)
    for text, fl in segs:
        r = p.add_run(text)
        r.font.size = Pt(9)
        r.italic = "i" in fl
        r.bold = "b" in fl
        if "s" in fl:
            r.font.subscript = True
    return p

def figure(path, width_in=6.3):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.keep_with_next = True   # caption stays with the figure
    p.add_run().add_picture(path, width=Inches(width_in))

# common inline fragments -------------------------------------------------
N  = ("N", "i");  D = ("D", "i");  X = ("X", "i");  T = ("T", "i")
r_ = ("r", "i")
PSI = ("Ψ", "i")
def psi_c():  return [("Ψ", "i"), ("c", "is")]
def kr():     return [("KR", "")]
def krn():    return [("KR", ""), ("norm", "s")]
def pr():     return [("PR", "")]

# ===================================================================== TITLE
h = doc.add_heading("", level=1)
h.add_run("S9.1 Kernel Rank")

# ---------------------------------------------------------------- definitions
body([("To characterise the effective dimensionality of the spin-wave "
       "reservoir feature space we construct, for ", ""), N,
      (" input samples, the feature matrix", "")])
equation_raw(EQ1_XML)

body([("where each row ", ""), ("ψ", "i"), ("i", "is"),
      (" is the flattened reservoir response of sample ", ""), ("i", "i"),
      (" over all spatial and temporal features, and ", ""), ("D", "i"),
      (" = ", ""), ("X", "i"), ("·", ""), ("T", "i"),
      (" is the number of spatial (", ""), X, (") by temporal (", ""), T,
      (") features per sample. After subtracting the column mean — so "
       "that the decomposition captures input-dependent variation rather than "
       "the static field — we take the singular value decomposition of "
       "the centred matrix ", "")] + psi_c() + [(",", "")])
equation(r"\Psi_c = U\,\Sigma\,V^\top")

body([("with singular values ", ""), ("σ", "i"), ("1", "is"),
      (" ≥ ", ""), ("σ", "i"), ("2", "is"),
      (" ≥ … ≥ 0. The cumulative explained variance retained by "
       "the leading ", ""), r_, (" components is", "")])
equation(r"C(r) = \frac{\sum_{i=1}^{r}\sigma_i^2}"
         r"{\sum_{i=1}^{\mathrm{min}(D,\,N-1)}\sigma_i^2}")

body([("and we define the kernel rank ", "")] + kr() +
     [(" as the smallest number of components whose cumulative variance "
       "reaches 95%,", "")])
equation(r"\mathrm{KR} = \mathrm{min}\{\,r : C(r) \geq 0.95\,\}")

body([("Because the centred matrix has at most ", ""), ("N", "i"),
      (" − 1 non-zero singular values, the achievable rank is bounded by "
       "min(", ""), D, (", ", ""), N, (" − 1); we report the normalised "
       "kernel rank", "")])
equation(r"\mathrm{KR}_{\mathrm{norm}} = \frac{\mathrm{KR}}{\mathrm{min}(D,\,N-1)}")

body([("In this work the feature dimension ", ""), ("D", "i"), (" = ", ""),
      ("X", "i"), ("·", ""), ("T", "i"),
      (" greatly exceeds the sample count ", ""), N,
      (" at every coarse-graining scale, so this bound is fixed at ", ""),
      ("N", "i"), (" − 1 = 255. To resolve how the variance is "
       "distributed across the retained dimensions we additionally report the "
       "participation ratio,", "")])
equation(r"\mathrm{PR} = \frac{\left(\sum_i \sigma_i^2\right)^2}"
         r"{\sum_i \sigma_i^4}")

body([("a continuous, variance-weighted effective rank: ", "")] + pr() +
     [(" is large when the variance is spread evenly across many modes and "
       "small when it is concentrated in a few.", "")])

# ----------------------------------------------------------------- figure S9.1
figure(os.path.join(FIGDIR, "fig5_cg_montage.png"))
caption([("Figure S9.1. ", "b"),
         ("Reservoir output (the ", "i"), ("m", "i"), ("z", "is"),
         (" field of one representative input sample) area-resampled to each "
          "spatial coarse-graining grid, from 50×50 to 4×4. These "
          "are the coarse-grained feature resolutions analysed in this "
          "section.", "i")])

# ----------------------------------------------------------------- results
body([("We computed ", "")] + kr() + [(", ", "")] + krn() + [(" and ", "")] +
     pr() + [(" for each coarse-grained feature set of Supplementary Fig. 4 "
       "— spatial grids from 50×50 down to 4×4, retaining the "
       "full temporal trajectory (", ""), N, (" = 256). The results are "
       "summarised in Table S9.1 and Fig. S9.2. The kernel rank is remarkably "
       "stable with spatial resolution: ", "")] + kr() +
     [(" = 25–26 (", "")] + krn() + [(" ≈ 0.10) for the four finest "
       "grids (50×50–20×20), 25 at 10×10, and only sags "
       "to 19–22 (", "")] + krn() + [(" ≈ 0.075–0.086) under "
       "the most aggressive coarse-graining (6×6 and 4×4). Thus 95% "
       "of the across-sample variance is carried by only ≈ 20–26 "
       "modes irrespective of spatial resolution — the input-"
       "discriminating variance is concentrated in a small set of large-scale "
       "spatial modes that survive coarse-graining.", "")])

# ----------------------------------------------------------------- table S9.1
tcap = doc.add_paragraph()
tcap.paragraph_format.space_before = Pt(8)
for text, fl in [("Table S9.1. ", "b"),
                  ("Kernel rank, normalised kernel rank and participation "
                   "ratio of the reservoir feature matrix across spatial "
                   "coarse-graining scales; the numerical rank of the centred "
                   "matrix is 255 = ", "i"), ("N", "ii"),
                  (" − 1 at every scale.", "i")]:
    r = tcap.add_run(text)
    r.font.size = Pt(9)
    r.bold = "b" in fl
    r.italic = "i" in fl

rows = [
    ("Grid", "D = X·T", "KR", "KR_norm", "PR", "rank(Ψc)"),
    ("50×50", "502,500", "26", "0.102", "24.3", "255"),
    ("40×40", "321,600", "26", "0.102", "24.1", "255"),
    ("30×30", "180,900", "26", "0.102", "24.1", "255"),
    ("20×20", "80,400",  "26", "0.102", "23.7", "255"),
    ("10×10", "20,100",  "25", "0.098", "20.6", "255"),
    ("8×8",   "12,864",  "23", "0.090", "17.2", "255"),
    ("6×6",   "7,236",   "19", "0.075", "13.4", "255"),
    ("4×4",   "3,216",   "22", "0.086", "8.4",  "255"),
]
tbl = doc.add_table(rows=len(rows), cols=6)
tbl.style = "Table Grid"
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
for ri, rowvals in enumerate(rows):
    for ci, val in enumerate(rowvals):
        cell = tbl.rows[ri].cells[ci]
        cell.text = ""
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(val)
        run.font.size = Pt(9.5)
        if ri == 0:
            run.bold = True

# ----------------------------------------------------------------- discussion
body([("This concentration of variance does not, however, imply a low-rank "
       "feature space. The centred matrix ", "")] + psi_c() +
     [(" is full rank — rank 255 = ", ""), ("N", "i"),
      (" − 1, the maximum — at every coarse-graining scale: the "
       "reservoir embeds all 256 inputs into linearly independent directions "
       "even after a 12-fold spatial downsampling, confirming a high-quality, "
       "separable feature map.", "")])

body([("The participation ratio resolves where coarse-graining begins to "
       "degrade the representation. ", "")] + pr() +
     [(" is flat at ≈ 24 for grids 50×50–20×20, then knees "
       "sharply at 10×10 (", "")] + pr() + [(" = 20.6) and falls "
       "monotonically to 8.4 at 4×4 (Fig. S9.2a). This knee at ≈ "
       "10×10 marks the spatial coarse-graining scale beyond which the "
       "fine-scale spin-wave structure can no longer be resolved — the "
       "spatial-Nyquist limit of the field — collapsing the variance "
       "onto progressively fewer large-scale modes. The reservoir therefore "
       "maintains a full-rank, input-separable feature space across a broad "
       "range of spatial resolutions, with a well-defined spatial-Nyquist "
       "limit revealed only by the variance-weighted rank.", "")])

# ----------------------------------------------------------------- figure S9.2
figure(os.path.join(FIGDIR, "figS91_kr_pr.png"))
caption([("Figure S9.2. ", "b"),
         ("(a) Kernel rank KR (= d95) and participation ratio PR of the "
          "feature matrix versus spatial coarse-graining grid; the numerical "
          "rank of Ψ", "i"), ("c", "is"),
         (" is 255 at every grid, and PR knees within the spatial-Nyquist "
          "band (shaded). (b) Cumulative explained variance ", "i"),
         ("C", "ii"), ("(", "i"), ("r", "ii"),
         (") for each grid; KR is the smallest ", "i"), ("r", "ii"),
         (" at which ", "i"), ("C", "ii"), ("(", "i"), ("r", "ii"),
         (") reaches 0.95 (dashed line).", "i")])

# ===================================================================== S9.1.1
sub = doc.add_heading("", level=2)
sub.add_run("S9.1.1 Spatial rank")

body([("The kernel rank above is computed on the full spatiotemporal volume, "
       "for which the feature dimension ", ""), ("D", "i"), (" = ", ""),
      ("X", "i"), ("·", ""), ("T", "i"), (" greatly exceeds ", ""), N,
      (" at every grid, so the rank ceiling is always ", ""), ("N", "i"),
      (" − 1. To isolate the purely spatial rank we repeat the analysis "
       "frame by frame: each of the ", ""), ("T", "i"),
      (" = 201 frames defines its own feature matrix ", ""),
      ("Ψ", "i"), ("(t)", "p"), (" ∈ ℝ", ""), ("N×X", "p"),
      (", whose rows are the coarse-grained spatial field of each sample at "
       "frame ", ""), ("t", "i"),
      ("; its feature dimension is just the spatial count ", ""),
      ("X", "i"), (" = ", ""), ("g", "i"), ("2", "p"),
      (". We report, for each grid, the frame of highest kernel rank,", "")])
equation(r"\mathrm{KR}^{\mathrm{spatial}} = \mathrm{max}_{t}\;"
         r"\mathrm{KR}\left(\Psi^{(t)}\right)")
body([("For the coarse grids (", ""), ("g", "i"), (" ≤ 15) the spatial "
       "feature dimension ", ""), ("X", "i"), (" = ", ""), ("g", "i"),
      ("2", "p"), (" falls below ", ""), N, (", so the rank ceiling becomes "
       "min(", ""), ("g", "i"), ("2", "p"), (", ", ""), ("N", "i"),
      (" − 1) = ", ""), ("g", "i"), ("2", "p"),
      (" — the analysis then tests directly whether the ", ""), ("g", "i"),
      ("2", "p"), (" spatial pixels of a single coarse-grained frame remain "
       "linearly independent across the ", ""), N, (" samples.", "")])

# ---- Table S9.2
tcap2 = doc.add_paragraph()
tcap2.paragraph_format.space_before = Pt(8)
for text, fl in [("Table S9.2. ", "b"),
                  ("Spatial (per-frame) kernel rank: for each grid the frame "
                   "t* of highest rank is reported. The numerical rank equals "
                   "the ceiling min(g², N − 1) at every grid.", "i")]:
    rr = tcap2.add_run(text)
    rr.font.size = Pt(9)
    rr.bold = "b" in fl
    rr.italic = "i" in fl

rows2 = [
    ("Grid", "X = g²", "best frame t*", "KR (d95)", "KR_norm", "numerical rank"),
    ("50×50", "2,500", "15",  "65", "0.255", "255"),
    ("40×40", "1,600", "19",  "51", "0.200", "255"),
    ("30×30", "900",   "19",  "46", "0.180", "255"),
    ("20×20", "400",   "18",  "37", "0.145", "255"),
    ("10×10", "100",   "116", "21", "0.210", "100"),
    ("8×8",   "64",    "116", "18", "0.281", "64"),
    ("6×6",   "36",    "123", "13", "0.361", "36"),
    ("4×4",   "16",    "136", "9",  "0.562", "16"),
]
tbl2 = doc.add_table(rows=len(rows2), cols=6)
tbl2.style = "Table Grid"
tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER
for ri, rowvals in enumerate(rows2):
    for ci, val in enumerate(rowvals):
        cell = tbl2.rows[ri].cells[ci]
        cell.text = ""
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(val)
        run.font.size = Pt(9.5)
        if ri == 0:
            run.bold = True

body([("Three features stand out. ", ""), ("First", "b"),
      (", the numerical rank of every per-frame matrix equals its ceiling "
       "min(", ""), ("g", "i"), ("2", "p"), (", ", ""), ("N", "i"),
      (" − 1); for the coarse grids this is exactly ", ""), ("g", "i"),
      ("2", "p"), (" (100, 64, 36 and 16 at 10×10–4×4), so the ", ""),
      ("g", "i"), ("2", "p"), (" spatial pixels of a single coarse-grained "
       "frame remain linearly independent across all 256 samples — the "
       "spatial feature map is non-degenerate at every resolution, with no "
       "rank collapse even at 4×4. ", ""), ("Second", "b"),
      (", the spatial kernel rank ", ""), ("KR", ""),
      (" falls monotonically with coarsening (65 → 9), in contrast to the "
       "near-constant full-volume ", ""), ("KR", ""),
      (" (≈ 25): a single sharp-wavefront frame spreads its variance across "
       "many more spatial modes than the temporally pooled volume. ", ""),
      ("Third", "b"), (", the highest-rank frame shifts systematically from "
       "early to late as the grid coarsens — ", ""), ("t", "i"), ("*", "ip"),
      (" ≈ 15–19 (≈ 0.3 ns, the early wavefront transient) for the fine "
       "grids, jumping to ", ""), ("t", "i"), ("*", "ip"),
      (" ≈ 116–136 (≈ 2.3–2.7 ns, the late ringdown) for the coarse grids "
       "(Fig. S9.3a). Fine grids resolve the fine-scale early wavefronts; "
       "coarse grids cannot, but do resolve the large-scale standing-wave "
       "pattern of the late ringdown. The rank-optimal readout time is "
       "therefore itself resolution-dependent.", "")])

figure(os.path.join(FIGDIR, "figS92_spatial_rank.png"))
caption([("Figure S9.3. ", "b"),
         ("(a) Per-frame spatial kernel rank ", "i"), ("d", "ii"),
         ("95 versus frame ", "i"), ("t", "ii"),
         (" for each grid; the peak (circle) shifts from the early wavefront "
          "transient (", "i"), ("t", "ii"),
          (" ≈ 15–19) at fine resolution to the late ringdown (", "i"),
         ("t", "ii"), (" ≈ 116–136) at coarse resolution. (b) Best-frame "
          "spatial kernel rank versus grid (navy), compared with the "
          "full-volume KR (crimson) and the rank ceiling min(g², N − 1) "
          "(dotted).", "i")])

body([("Because the optimal frame differs between grids, the entries of "
       "Table S9.2 are not a fixed-time comparison — each reports the best "
       "spatial frame for its own resolution.", "")])

doc.save(OUT)
print("saved", OUT)
